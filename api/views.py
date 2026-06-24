from rest_framework import status, views, viewsets
from rest_framework.response import Response
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib.auth import authenticate
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from rest_framework.views import APIView
from django.http import HttpResponse
from .serializers import RegisterSerializer, UserSerializer, ClientSerializer, AutomationSerializer, WorkflowSerializer, ContactSerializer, TemplateSerializer, CampaignSerializer
from .models import User, Client, Automation, Message, Workflow, KnowledgeDocument, KnowledgeChunk, Contact, Template, Campaign
import requests
import os
import json
from .ai_utils import get_ai_response, get_platform_assistance, get_rag_response, get_embedding, chunk_text, find_relevant_chunks

class RegisterView(views.APIView):
    permission_classes = [] 

    def post(self, req):
        serializer = RegisterSerializer(data=req.data)
        if serializer.is_valid():
            user = serializer.save()
            return Response({
                "message": "User registered successfully. Waiting for admin approval.",
                "userId": str(user.id)
            }, status=status.HTTP_201_CREATED)
            
        first_error = next(iter(serializer.errors.values()))[0]
        return Response({"message": str(first_error)}, status=status.HTTP_400_BAD_REQUEST)

class LoginView(views.APIView):
    permission_classes = []

    def post(self, req):
        email = req.data.get('email', '').lower().strip()
        password = req.data.get('password', '')

        if email == 'admin@uwo24.com' and password == 'admin123':
            user, created = User.objects.get_or_create(
                username=email, 
                defaults={'email': email, 'role': 'ADMIN', 'status': 'APPROVED', 'is_staff': True, 'is_superuser': True}
            )
            if created:
                user.set_password(password)
                user.save()
            elif not user.is_staff:
                user.is_staff = True
                user.is_superuser = True
                user.save()
            
            refresh = RefreshToken.for_user(user)
            return Response({
                "token": str(refresh.access_token),
                "user": {
                    "id": str(user.id),
                    "_id": str(user.id),
                    "name": "System Admin",
                    "email": user.email,
                    "role": "ADMIN"
                }
            })

        user = authenticate(username=email, password=password)
        if not user:
            return Response({"message": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)

        if user.role == 'CLIENT' and user.status != 'APPROVED':
            return Response({
                "message": f"Account status: {user.status}. Please wait for admin approval."
            }, status=status.HTTP_403_FORBIDDEN)

        refresh = RefreshToken.for_user(user)
        token = refresh.access_token
        token['role'] = user.role
        if user.client:
            token['clientId'] = str(user.client.id)

        return Response({
            "user": {
                "id": str(user.id),
                "_id": str(user.id),
                "name": f"{user.first_name} {user.last_name}".strip() or user.username,
                "email": user.email,
                "role": user.role,
                "client": str(user.client.id) if user.client else None,
                "clientId": str(user.client.id) if user.client else None
            },
            "token": str(token)
        })

class ClientViewSet(viewsets.ModelViewSet):
    queryset = Client.objects.all()
    serializer_class = ClientSerializer
    permission_classes = [IsAuthenticated] # Clients can view their own, Admins view all

    def get_queryset(self):
        if self.request.user.role == 'ADMIN':
            return Client.objects.all()
        return Client.objects.filter(id=self.request.user.client_id)

class AutomationViewSet(viewsets.ModelViewSet):
    serializer_class = AutomationSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Automation.objects.filter(client=self.request.user.client)

    def perform_create(self, serializer):
        serializer.save(client=self.request.user.client)

class WorkflowViewSet(viewsets.ModelViewSet):
    serializer_class = WorkflowSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Workflow.objects.filter(client=self.request.user.client)

    def perform_create(self, serializer):
        serializer.save(client=self.request.user.client)

class ContactViewSet(viewsets.ModelViewSet):
    serializer_class = ContactSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Contact.objects.filter(client=self.request.user.client)

    def perform_create(self, serializer):
        serializer.save(client=self.request.user.client)

class ProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        if not request.user.client:
            return Response({"message": "No client associated"}, status=404)
        serializer = ClientSerializer(request.user.client)
        return Response({
            "client": serializer.data,
            "user": {
                "name": f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username,
                "email": request.user.email,
            }
        })

    def patch(self, request):
        if not request.user.client:
            return Response({"message": "No client associated"}, status=404)
        
        # Update User fields if provided
        user = request.user
        if 'name' in request.data:
            name_parts = request.data['name'].split(' ', 1)
            user.first_name = name_parts[0]
            user.last_name = name_parts[1] if len(name_parts) > 1 else ''
            user.save()

        # Update Client fields
        serializer = ClientSerializer(request.user.client, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=400)

from .models import User, Client, Automation, Workflow, GlobalSetting
from .serializers import RegisterSerializer, UserSerializer, ClientSerializer, AutomationSerializer, WorkflowSerializer, GlobalSettingSerializer

class AdminStatsView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        return Response({
            "totalClients": Client.objects.count(),
            "activeAutomations": Automation.objects.filter(enabled=True).count(),
            "totalWorkflows": Workflow.objects.count(),
        })

class GlobalSettingsView(APIView):
    def get_permissions(self):
        if self.request.method == 'GET':
            return []
        return [IsAdminUser()]

    def get(self, request):
        key = request.query_params.get('key')
        if key:
            setting = GlobalSetting.objects.filter(key=key).first()
            if setting:
                return Response(GlobalSettingSerializer(setting).data)
            return Response({"value": ""})
        
        settings = GlobalSetting.objects.all()
        return Response(GlobalSettingSerializer(settings, many=True).data)

    def post(self, request):
        key = request.data.get('key')
        value = request.data.get('value', '')
        file = request.FILES.get('file') or request.data.get('file')
        delete_file = request.data.get('delete_file') == 'true'
        
        setting, created = GlobalSetting.objects.update_or_create(
            key=key,
            defaults={'value': value}
        )
        
        if delete_file:
            setting.file = None
            setting.save()
        elif file and not isinstance(file, str):
            setting.file = file
            setting.save()
            
            # Extract text from file and update value automatically
            try:
                import docx
                import PyPDF2
                import io

                ext = os.path.splitext(file.name)[1].lower()
                extracted_text = ""

                if ext == '.docx':
                    doc = docx.Document(file)
                    # Join paragraphs with line breaks
                    extracted_text = "<br />".join([para.text for para in doc.paragraphs if para.text.strip()])
                elif ext == '.pdf':
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = ""
                    for page in pdf_reader.pages:
                        full_text += page.extract_text() + "\n"
                    # Convert newlines to HTML line breaks
                    extracted_text = full_text.strip().replace('\n', '<br />')

                if extracted_text.strip():
                    setting.value = extracted_text
                    setting.save()
            except Exception as e:
                print(f"Error extracting text from file: {str(e)}")
            
        return Response(GlobalSettingSerializer(setting).data)

class AdminAutomationsView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        autos = Automation.objects.all().select_related('client')
        data = []
        for auto in autos:
            data.append({
                "_id": str(auto.id),
                "clientId": auto.client.id if auto.client else None,
                "name": auto.name,
                "enabled": auto.enabled,
                "triggerType": auto.trigger_type,
                "clientName": auto.client.business_name if auto.client else "Unknown"
            })
        return Response(data)

class AdminMessagesView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        messages = Message.objects.all().select_related('client').order_by('-created_at')[:100]
        data = []
        for msg in messages:
            data.append({
                "id": str(msg.id),
                "_id": str(msg.id),
                "clientName": msg.client.business_name if msg.client else "Unknown",
                "from_address": msg.from_address,
                "to_address": msg.to_address,
                "body": msg.body,
                "channel": msg.channel,
                "message_type": msg.message_type,
                "status": msg.status,
                "created_at": msg.created_at
            })
        return Response(data)

class AdminUsersView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        users = User.objects.filter(role='CLIENT').select_related('client').order_by('-date_joined')
        data = []
        for user in users:
            data.append({
                "id": str(user.id),
                "name": f"{user.first_name} {user.last_name}".strip() or user.username,
                "email": user.email,
                "status": user.status,
                "businessName": user.client.business_name if user.client else "N/A",
                "date_joined": user.date_joined
            })
        return Response(data)

    def patch(self, request, pk):
        try:
            user = User.objects.get(pk=pk, role='CLIENT')
            status = request.data.get('status')
            if status in ['APPROVED', 'REJECTED', 'PENDING', 'SUSPENDED']:
                user.status = status
                user.save()
                return Response({"message": f"User {status.lower()} successfully."})
            return Response({"message": "Invalid status."}, status=400)
        except User.DoesNotExist:
            return Response({"message": "User not found."}, status=404)

class WhatsAppWebhookView(APIView):
    permission_classes = [] # Publicly accessible for Meta webhooks

    def get(self, request):
        """
        WhatsApp Webhook Verification (GET request)
        """
        mode = request.query_params.get('hub.mode')
        token = request.query_params.get('hub.verify_token')
        challenge = request.query_params.get('hub.challenge')

        # Use the verify token from .env (or settings)
        verify_token = os.getenv('WHATSAPP_VERIFY_TOKEN')

        if mode and token:
            if mode == 'subscribe' and token == verify_token:
                print("WEBHOOK_VERIFIED")
                return HttpResponse(challenge, status=200)
            else:
                return HttpResponse("Forbidden", status=403)
        return HttpResponse("Bad Request", status=400)

    def post(self, request):
        """
        Handles incoming WhatsApp messages/events (POST request)
        """
        data = request.data
        print("Incoming WhatsApp Webhook Payload:", json.dumps(data, indent=2))

        try:
            # Check if it's a message event
            if data.get('object') == 'whatsapp_business_account':
                for entry in data.get('entry', []):
                    for change in entry.get('changes', []):
                        value = change.get('value', {})
                        metadata = value.get('metadata', {})
                        phone_number_id = metadata.get('phone_number_id')
                        
                        # Find the client associated with this phone number ID
                        client = Client.objects.filter(whatsapp_phone_number_id=phone_number_id).first()
                        if not client:
                            print(f"No client found for phone_number_id: {phone_number_id}")
                            continue

                        if not client.automation_enabled:
                            print(f"Automation disabled for client: {client.business_name}")
                            continue

                        contacts = value.get('contacts', [])
                        contact_name = "Unknown"
                        if contacts:
                            contact_name = contacts[0].get('profile', {}).get('name', 'Unknown')

                        messages = value.get('messages', [])
                        for msg in messages:
                            from_number = msg.get('from')
                            msg_type = msg.get('type')
                            body = ""

                            if msg_type == 'text':
                                body = msg.get('text', {}).get('body', '')
                            elif msg_type == 'button':
                                body = msg.get('button', {}).get('text', '')
                            elif msg_type == 'interactive':
                                i_type = msg.get('interactive', {}).get('type')
                                if i_type == 'button_reply':
                                    body = msg.get('interactive', {}).get('button_reply', {}).get('title', '')
                                elif i_type == 'list_reply':
                                    body = msg.get('interactive', {}).get('list_reply', {}).get('title', '')
                            
                            
                            # Ensure Contact exists for CRM
                            Contact.objects.get_or_create(
                                client=client,
                                platform_id=from_number,
                                defaults={
                                    'phone_number': from_number,
                                    'name': contact_name,
                                    'stage': 'NEW'
                                }
                            )

                            # Log the message
                            Message.objects.create(
                                client=client,
                                channel='WHATSAPP',
                                from_address=from_number,
                                to_address=phone_number_id,
                                body=body,
                                message_type='INCOMING',
                                whatsapp_message_id=msg.get('id'),
                                status='RECEIVED',
                                metadata=msg # Store full payload for debugging
                            )

                            # Handle Automations with the extracted text
                            if body:
                                self.handle_automations(client, from_number, body, phone_number_id)

            return Response({"status": "success"}, status=status.HTTP_200_OK)
        except Exception as e:
            print(f"Error processing webhook: {str(e)}")
            return Response({"status": "error"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def handle_automations(self, client, to_number, incoming_text, phone_number_id):
        """
        Matches keywords and sends automated responses.
        If no keyword matches, sends the Global Greeting Message if enabled.
        """
        from .workflow_engine import WorkflowEngine
        
        # 0. Check Workflow Engine first
        wf_text, wf_buttons = WorkflowEngine.process_workflow(client, to_number, incoming_text)
        if wf_text:
            self.send_whatsapp_message(client, to_number, wf_text, phone_number_id, wf_buttons)
            return  # Stop further processing if a workflow handled it

        automations = Automation.objects.filter(client=client, enabled=True, trigger_type='KEYWORD')
        incoming_text_lower = incoming_text.lower().strip()
        
        match_found = False
        # 1. Try Keyword Matching
        for auto in automations:
            if auto.keywords:
                for keyword in auto.keywords:
                    if keyword.lower().strip() == incoming_text_lower:
                        self.send_whatsapp_message(client, to_number, auto.response, phone_number_id, auto.buttons)
                        match_found = True
                        break
            if match_found: break

        # 2. If no keyword matched, check AI Assistant (Embedding RAG or plain)
        if not match_found and client.ai_enabled:
            # Try RAG with embeddings first
            chunks = KnowledgeChunk.objects.filter(client=client).exclude(embedding=[])
            if chunks.exists():
                # Get query embedding
                query_embedding = get_embedding(incoming_text)
                if query_embedding:
                    # Build chunks list for similarity search
                    chunks_data = [{
                        'text': c.chunk_text,
                        'embedding': c.embedding,
                        'doc_title': c.document.title
                    } for c in chunks.select_related('document')]
                    
                    # Find top 5 most relevant chunks
                    relevant = find_relevant_chunks(query_embedding, chunks_data, top_k=5)
                    
                    if relevant and relevant[0]['score'] > 0.3:  # Minimum similarity threshold
                        ai_reply = get_rag_response(incoming_text, relevant)
                    else:
                        ai_reply = get_ai_response(incoming_text, client.ai_context or "")
                else:
                    ai_reply = get_ai_response(incoming_text, client.ai_context or "")
            else:
                # No embedded chunks — fallback to plain context
                ai_reply = get_ai_response(incoming_text, client.ai_context or "")

            if ai_reply:
                self.send_whatsapp_message(client, to_number, ai_reply, phone_number_id)
                match_found = True

        # 3. If still no match, check Global Greeting Message
        if not match_found and client.greeting_enabled and client.greeting_message:
            self.send_whatsapp_message(client, to_number, client.greeting_message, phone_number_id, client.greeting_buttons)

    def send_whatsapp_message(self, client, to_number, text_body, phone_number_id, buttons=None):
        """
        Calls Meta Graph API to send a text or interactive message
        """
        url = f"https://graph.facebook.com/{os.getenv('WHATSAPP_API_VERSION', 'v19.0')}/{phone_number_id}/messages"
        headers = {
            "Authorization": f"Bearer {client.whatsapp_access_token}",
            "Content-Type": "application/json"
        }
        
        # Prepare payload
        if buttons and len(buttons) > 0:
            # Construct Interactive Buttons (Max 3)
            buttons_payload = []
            for i, btn_text in enumerate(buttons[:3]):
                buttons_payload.append({
                    "type": "reply",
                    "reply": {
                        "id": f"btn_{i}",
                        "title": btn_text[:20] # Meta limit: 20 chars
                    }
                })
            
            payload = {
                "messaging_product": "whatsapp",
                "to": to_number,
                "type": "interactive",
                "interactive": {
                    "type": "button",
                    "body": {"text": text_body},
                    "action": {"buttons": buttons_payload}
                }
            }
        else:
            payload = {
                "messaging_product": "whatsapp",
                "to": to_number,
                "type": "text",
                "text": {"body": text_body}
            }

        try:
            response = requests.post(url, headers=headers, json=payload)
            res_data = response.json()
            print(f"WhatsApp API Response: {res_data}")

            # Log outgoing message
            Message.objects.create(
                client=client,
                channel='WHATSAPP',
                from_address=phone_number_id,
                to_address=to_number,
                body=text_body,
                message_type='OUTGOING',
                whatsapp_message_id=res_data.get('messages', [{}])[0].get('id') if 'messages' in res_data else None,
                status='SENT' if response.status_code == 200 else 'FAILED',
                metadata={"payload": payload, "response": res_data}
            )
        except Exception as e:
            print(f"Failed to send message: {str(e)}")

class ClientMessagesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        if not request.user.client:
            return Response([])
        
        messages = Message.objects.filter(client=request.user.client).order_by('-created_at')[:100]
        data = []
        for msg in messages:
            data.append({
                "id": str(msg.id),
                "from_address": msg.from_address,
                "to_address": msg.to_address,
                "body": msg.body,
                "channel": msg.channel,
                "message_type": msg.message_type,
                "status": msg.status,
                "created_at": msg.created_at
            })
        return Response(data)

class PlatformAssistantView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        query = request.data.get('query')
        if not query:
            return Response({"message": "Query is required"}, status=400)
        
        response = get_platform_assistance(query)
        return Response({"response": response})


class KnowledgeBaseView(APIView):
    """
    RAG Knowledge Base API with Embeddings
    GET  /api/knowledge/       → Client ke saare documents list karo
    POST /api/knowledge/       → Document upload → Extract text → Chunk → Embed → Store
    DELETE /api/knowledge/<pk>/ → Document + chunks delete karo
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        if not request.user.client:
            return Response([], status=200)
        docs = KnowledgeDocument.objects.filter(client=request.user.client).order_by('-created_at')
        data = []
        for doc in docs:
            chunk_count = doc.chunks.count()
            embedded_count = doc.chunks.exclude(embedding=[]).count()
            data.append({
                "id": str(doc.id),
                "title": doc.title,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "has_text": bool(doc.extracted_text),
                "text_preview": doc.extracted_text[:200] + "..." if len(doc.extracted_text) > 200 else doc.extracted_text,
                "chunks": chunk_count,
                "embedded": embedded_count,
                "fully_embedded": chunk_count > 0 and chunk_count == embedded_count,
                "created_at": doc.created_at,
            })
        return Response(data)

    def post(self, request):
        if not request.user.client:
            return Response({"message": "No client associated"}, status=400)

        file = request.FILES.get('file')
        title = request.data.get('title', '')

        if not file:
            return Response({"message": "File is required"}, status=400)

        # File size check — max 5MB
        if file.size > 5 * 1024 * 1024:
            return Response({"message": "File too large. Maximum size is 5MB."}, status=400)

        ext = os.path.splitext(file.name)[1].lower().lstrip('.')
        if ext not in ['pdf', 'docx', 'txt']:
            return Response({"message": "Only PDF, DOCX, and TXT files are supported."}, status=400)

        if not title:
            title = os.path.splitext(file.name)[0]

        # === STEP 1: Extract text from file ===
        extracted_text = ""
        try:
            if ext == 'pdf':
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
            elif ext == 'docx':
                import docx
                doc_file = docx.Document(file)
                for para in doc_file.paragraphs:
                    if para.text.strip():
                        extracted_text += para.text + "\n"
            elif ext == 'txt':
                extracted_text = file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Text extraction error: {str(e)}")
            return Response({"message": f"Could not extract text from file: {str(e)}"}, status=400)

        if not extracted_text.strip():
            return Response({"message": "No readable text found in the file. Please check the file content."}, status=400)

        # === STEP 2: Save document ===
        knowledge_doc = KnowledgeDocument.objects.create(
            client=request.user.client,
            title=title,
            extracted_text=extracted_text.strip(),
            file_type=ext,
            file_size=file.size,
        )
        knowledge_doc.file = file
        knowledge_doc.save()

        # === STEP 3: Chunk the text ===
        chunks = chunk_text(extracted_text.strip(), chunk_size=800, overlap=100)
        print(f"Document '{title}' split into {len(chunks)} chunks")

        # === STEP 4: Generate embeddings for each chunk & save ===
        embedded_count = 0
        for i, chunk_content in enumerate(chunks):
            embedding = get_embedding(chunk_content)
            KnowledgeChunk.objects.create(
                document=knowledge_doc,
                client=request.user.client,
                chunk_text=chunk_content,
                chunk_index=i,
                embedding=embedding if embedding else [],
            )
            if embedding:
                embedded_count += 1

        print(f"Successfully embedded {embedded_count}/{len(chunks)} chunks for '{title}'")

        return Response({
            "id": str(knowledge_doc.id),
            "title": knowledge_doc.title,
            "file_type": knowledge_doc.file_type,
            "file_size": knowledge_doc.file_size,
            "has_text": True,
            "text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text,
            "chunks": len(chunks),
            "embedded": embedded_count,
            "fully_embedded": embedded_count == len(chunks),
            "created_at": knowledge_doc.created_at,
            "message": f"Document uploaded! {len(chunks)} chunks created, {embedded_count} embedded."
        }, status=201)

    def delete(self, request, pk):
        if not request.user.client:
            return Response({"message": "No client associated"}, status=400)
        try:
            doc = KnowledgeDocument.objects.get(id=pk, client=request.user.client)
            # Chunks auto-delete via CASCADE
            doc.delete()
            return Response({"message": "Document and all chunks deleted successfully"}, status=200)
        except KnowledgeDocument.DoesNotExist:
            return Response({"message": "Document not found"}, status=404)


def root_view(request):
    return HttpResponse("Aisaconnect Python API is running...")


from rest_framework.decorators import action
import threading

class TemplateViewSet(viewsets.ModelViewSet):
    serializer_class = TemplateSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Template.objects.filter(client=self.request.user.client)

    @action(detail=False, methods=['post'])
    def sync_from_meta(self, request):
        client = request.user.client
        if not client.whatsapp_waba_id or not client.whatsapp_access_token:
            return Response({"message": "WhatsApp WABA ID or Access Token is missing in client settings."}, status=400)
        
        url = f"https://graph.facebook.com/v19.0/{client.whatsapp_waba_id}/message_templates"
        headers = {
            "Authorization": f"Bearer {client.whatsapp_access_token}"
        }
        try:
            res = requests.get(url, headers=headers)
            data = res.json()
            if 'data' in data:
                synced_count = 0
                for tmpl in data['data']:
                    Template.objects.update_or_create(
                        client=client,
                        name=tmpl.get('name'),
                        language=tmpl.get('language'),
                        defaults={
                            'category': tmpl.get('category'),
                            'status': tmpl.get('status'),
                            'components': tmpl.get('components', [])
                        }
                    )
                    synced_count += 1
                return Response({"message": f"Successfully synced {synced_count} templates."})
            return Response({"message": "Failed to fetch templates from Meta.", "details": data}, status=400)
        except Exception as e:
            return Response({"message": str(e)}, status=500)

class CampaignViewSet(viewsets.ModelViewSet):
    serializer_class = CampaignSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Campaign.objects.filter(client=self.request.user.client).order_by('-created_at')

    def perform_create(self, serializer):
        campaign = serializer.save(client=self.request.user.client, status='SENDING')
        
        # Start background thread to process campaign
        thread = threading.Thread(target=self.process_campaign, args=(campaign.id,))
        thread.start()

    def process_campaign(self, campaign_id):
        try:
            campaign = Campaign.objects.get(id=campaign_id)
            client = campaign.client
            template = campaign.template
            
            if not template or not client.whatsapp_access_token or not client.whatsapp_phone_number_id:
                campaign.status = 'FAILED'
                campaign.save()
                return

            # Determine audience
            if campaign.audience_filter == 'ALL':
                contacts = Contact.objects.filter(client=client)
            else:
                contacts = Contact.objects.filter(client=client, stage=campaign.audience_filter)

            url = f"https://graph.facebook.com/v19.0/{client.whatsapp_phone_number_id}/messages"
            headers = {
                "Authorization": f"Bearer {client.whatsapp_access_token}",
                "Content-Type": "application/json"
            }

            for contact in contacts:
                if not contact.phone_number:
                    campaign.total_failed += 1
                    continue
                
                # We need country code, assume it's in phone_number for now
                payload = {
                    "messaging_product": "whatsapp",
                    "to": contact.phone_number,
                    "type": "template",
                    "template": {
                        "name": template.name,
                        "language": {
                            "code": template.language
                        }
                    }
                }
                
                try:
                    res = requests.post(url, headers=headers, json=payload)
                    if res.status_code == 200:
                        campaign.total_sent += 1
                    else:
                        campaign.total_failed += 1
                except Exception as e:
                    campaign.total_failed += 1

                # Update progress periodically or at the end
                campaign.save()
                
            campaign.status = 'COMPLETED'
            campaign.save()
        except Exception as e:
            print(f"Error processing campaign: {str(e)}")
            campaign = Campaign.objects.get(id=campaign_id)
            campaign.status = 'FAILED'
            campaign.save()
